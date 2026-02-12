# atp_docx_insert.py
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from PIL import Image
import io
import os


TEXT_PLACEHOLDERS = [
    "[SK_1]",
    "[SITE_ID1]",
    "[SITE_NAME1]",
    "[SITE_ID]",
    "[SITE_NAME]",
    "[HOSTNAME]",
    "[SCOPE_OF_WORK]",
    "[DEVICE_TYPE]",
    "[PROJECT_CODE]",
    "[DATE]",
    "[ENGINEER]",
    "[LOCATION]",
    "[ADDRESS]",
]


class ATPDocxInserter:
    """
    Handles photo insertion and text replacement in DOCX ATP templates.
    """

    def __init__(self, docx_path):
        """
        Initialize with a DOCX template.

        Args:
            docx_path: Path to the DOCX template file
        """
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        self.photo_mappings = self.detect_photo_placeholders()
        self.text_mappings = self.detect_text_placeholders()

    def detect_photo_placeholders(self):
        """
        Detect photo placeholder text in the DOCX document.
        Looks for patterns like [PHOTO_FRONT_VIEW] in paragraphs and tables.
        """
        mappings = []

        for para_idx, paragraph in enumerate(self.doc.paragraphs):
            text = paragraph.text.strip()
            if self.is_photo_placeholder(text):
                mappings.append({
                    "type": "paragraph",
                    "paragraph_index": para_idx,
                    "placeholder": text,
                    "photo_type": self.get_photo_type(text),
                    "location": f"Paragraph {para_idx + 1}",
                })

        for table_idx, table in enumerate(self.doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        text = paragraph.text.strip()
                        if self.is_photo_placeholder(text):
                            mappings.append({
                                "type": "table_cell",
                                "table_index": table_idx,
                                "row_index": row_idx,
                                "cell_index": cell_idx,
                                "paragraph_index": para_idx,
                                "placeholder": text,
                                "photo_type": self.get_photo_type(text),
                                "location": f"Table {table_idx + 1}, Cell ({row_idx + 1},{cell_idx + 1})",
                            })

        return mappings

    # def detect_text_placeholders(self):
    #     """Detect text placeholders in the document."""
    #     mappings = []
    #     seen = set()

    #     def scan_text(text, location):
    #         for placeholder in TEXT_PLACEHOLDERS:
    #             if placeholder in text and placeholder not in seen:
    #                 seen.add(placeholder)
                    
    #                 # Keep the original format for display
    #                 display_name = placeholder.strip("[]").replace("_", " ").title()
                    
    #                 # Create a clean key - KEEP UNDERSCORES for proper matching
    #                 key = placeholder.strip("[]").lower()
                    
    #                 mappings.append({
    #                     "placeholder": placeholder,
    #                     "placeholder_key": key,  # Keep underscores (site_id, not siteid)
    #                     "display_name": display_name,
    #                     "location": location,
    #                     "required": key in ["site_id", "site_name", "project_code", 
    #                                       "sk_1", "site_id1", "site_name1"],
    #                 })

    #     # Scan paragraphs
    #     for i, paragraph in enumerate(self.doc.paragraphs):
    #         scan_text(paragraph.text, f"Paragraph {i + 1}")

    #     # Scan tables
    #     for t_i, table in enumerate(self.doc.tables):
    #         for r_i, row in enumerate(table.rows):
    #             for c_i, cell in enumerate(row.cells):
    #                 for p_i, paragraph in enumerate(cell.paragraphs):
    #                     scan_text(
    #                         paragraph.text,
    #                         f"Table {t_i + 1}, Cell ({r_i + 1},{c_i + 1})"
    #                     )

    #     return mappings


    def detect_text_placeholders(self):
         """Detect text placeholders in the document."""
         mappings = []
         seen = set()

         def scan_text(text, location):
             for placeholder in TEXT_PLACEHOLDERS:
                 if placeholder in text and placeholder not in seen:
                     seen.add(placeholder)

                     # Keep the original format
                     display_name = placeholder.strip("[]").replace("_", " ").title()

                     # Create a clean key - convert to lowercase
                     key = placeholder.strip("[]").lower()

                     # Determine if required
                     required = key in [
                         "site_id", "site_name", "project_code", 
                         "sk_1", "site_id1", "site_name1", "date"
                     ]

                     mappings.append({
                         "placeholder": placeholder,
                         "placeholder_key": key,  # 'date', 'engineer', etc.
                         "display_name": display_name,
                         "location": location,
                         "required": required,
                     })

         # Scan paragraphs
         for i, paragraph in enumerate(self.doc.paragraphs):
             scan_text(paragraph.text, f"Paragraph {i + 1}")

         # Scan tables
         for t_i, table in enumerate(self.doc.tables):
             for r_i, row in enumerate(table.rows):
                 for c_i, cell in enumerate(row.cells):
                     for p_i, paragraph in enumerate(cell.paragraphs):
                         scan_text(
                             paragraph.text,
                             f"Table {t_i + 1}, Cell ({r_i + 1},{c_i + 1})"
                         )

         return mappings
    

    def is_photo_placeholder(self, text):
        """Check if text is a photo placeholder."""
        photo_placeholders = [
            "[BEFORE_TOWER1]",
            "[AFTER_TOWER1]",
            "[PHOTO_FRONT_SPACE]",
            "[PHOTO_REAR_SPACE]",
            "[PHOTO_POWER_CABLE]",
            "[PHOTO_GROUNDING]",
            "[PHOTO_CONNECTION]",
            "[PHOTO]",
            "[IMAGE]",
            "[PICTURE]",
        ]
        return text.upper() in [p.upper() for p in photo_placeholders]

    def get_photo_type(self, placeholder):
        """Get human-readable photo type from placeholder."""
        photo_types = {
            "[BEFORE_TOWER1]": "Before Tower 1",
            "[AFTER_TOWER1]": "After Tower 1",
            "[PHOTO_FRONT_SPACE]": "Front Space",
            "[PHOTO_REAR_SPACE]": "Rear Space",
            "[PHOTO_POWER_CABLE]": "Power Cable",
            "[PHOTO_GROUNDING]": "Grounding Cable",
            "[PHOTO_CONNECTION]": "Connection Router",
            "[PHOTO]": "Photo",
            "[IMAGE]": "Image",
            "[PICTURE]": "Picture",
        }
        return photo_types.get(placeholder.upper(), placeholder)

    def insert_photo(
        self, mapping_index, photo_path, width_inches=3.0, height_inches=2.0
    ):
        """
        Insert a photo at the specified placeholder location.

        Args:
            mapping_index: Index in photo_mappings list
            photo_path: Path to the photo file
            width_inches: Width of image in inches
            height_inches: Height of image in inches
        """
        if mapping_index >= len(self.photo_mappings):
            return False

        mapping = self.photo_mappings[mapping_index]

        try:
            if mapping["type"] == "paragraph":
                # Insert photo in paragraph
                paragraph = self.doc.paragraphs[mapping["paragraph_index"]]
                # Clear the placeholder text
                paragraph.clear()
                # Add the image
                run = paragraph.add_run()
                run.add_picture(
                    photo_path, width=Inches(width_inches), height=Inches(height_inches)
                )
                # Center the image
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            elif mapping["type"] == "table_cell":
                # Insert photo in table cell
                table = self.doc.tables[mapping["table_index"]]
                cell = table.cell(mapping["row_index"], mapping["cell_index"])
                paragraph = cell.paragraphs[mapping["paragraph_index"]]
                # Clear the placeholder text
                paragraph.clear()
                # Add the image
                run = paragraph.add_run()
                run.add_picture(
                    photo_path, width=Inches(width_inches), height=Inches(height_inches)
                )
                # Center the image
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            return True

        except Exception as e:
            print(f"Error inserting photo: {e}")
            return False

    def replace_text(self, placeholder, replacement_text):
        """
        Replace a specific placeholder with text.
        
        Args:
            placeholder: The exact placeholder text to replace (e.g., '[SITE_ID]')
            replacement_text: The text to insert
            
        Returns:
            int: Number of replacements made
        """
        if not placeholder or replacement_text is None:
            return 0
            
        placeholder_upper = placeholder.upper()
        replacement_count = 0

        def replace_in_paragraph(paragraph):
            nonlocal replacement_count
            # If paragraph has no runs, create one
            if not paragraph.runs:
                paragraph.add_run()
            
            full_text = paragraph.text
            if placeholder_upper not in full_text.upper():
                return
            
            # Handle runs properly
            for run in paragraph.runs:
                if placeholder_upper in run.text.upper():
                    run.text = run.text.replace(placeholder, replacement_text)
                    replacement_count += 1

        # Process all paragraphs
        for paragraph in self.doc.paragraphs:
            replace_in_paragraph(paragraph)

        # Process all tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph)
        
        return replacement_count

    def replace_all_text(self, text_values):
        """
        Replace all text placeholders with values from dictionary.
        
        Args:
            text_values: Dictionary mapping placeholder keys to values
                        Example: {"site_id": "SITE-001", "sk_1": "ABC123"}
        
        Returns:
            dict: Count of replacements per key
        """
        replacements = {}
        
        for key, value in text_values.items():
            if not value or not isinstance(value, str):
                continue

            # Convert key to placeholder format
            # Handle both with and without underscores
            key_upper = key.upper()
            
            # Try different placeholder formats
            placeholders_to_try = [
                f"[{key_upper}]",  # [SITE_ID]
                f"[{key_upper.replace('_', '')}]",  # [SITEID]
            ]
            
            # Also try the original format from TEXT_PLACEHOLDERS
            for placeholder in TEXT_PLACEHOLDERS:
                if placeholder.strip('[]').lower() == key.lower():
                    placeholders_to_try.append(placeholder)
            
            # Try each placeholder format
            for placeholder in set(placeholders_to_try):  # Use set to remove duplicates
                count = self.replace_text(placeholder, value)
                if count > 0:
                    replacements[key] = replacements.get(key, 0) + count
                    break  # Stop after first successful replacement
        
        return replacements

    def get_available_photo_slots(self):
        """Get photo slots information for frontend display."""
        slots = []

        for idx, mapping in enumerate(self.photo_mappings):
            slots.append(
                {
                    "slot_index": idx,
                    "type": mapping["photo_type"],
                    "photo_type": mapping["photo_type"],
                    "location": mapping["location"],
                    "placeholder": mapping["placeholder"],
                    "status": "empty",
                }
            )

        return slots

    def get_available_text_fields(self):
        """Get text fields information for frontend display."""
        text_fields = []
        seen_placeholders = set()

        for mapping in self.text_mappings:
            placeholder = mapping["placeholder"]
            if placeholder not in seen_placeholders:
                seen_placeholders.add(placeholder)

                text_fields.append(
                    {
                        "placeholder": placeholder,
                        "placeholder_key": mapping["placeholder_key"],  # Keep the key with underscores
                        "display_name": mapping["display_name"],
                        "location": mapping["location"],
                        "required": mapping["required"],
                    }
                )

        return text_fields

    def save(self, output_path):
        """Save the modified document."""
        self.doc.save(output_path)